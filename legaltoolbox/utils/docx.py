import docx
from docx.shared import Pt
from docx.shared import Mm

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Cm

def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"

def add_hyperlink(paragraph, url, text):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_transposition_table() -> docx.Document:

    doc = docx.Document() 

    new_width, new_height = doc.sections[-1].page_height,  doc.sections[-1].page_width
    new_section = doc.sections[-1]
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
        
    new_section.top_margin = Cm(0.5)
    new_section.bottom_margin = Cm(0.5)
    new_section.left_margin = Cm(1)
    new_section.right_margin = Cm(1)

    doc.add_heading(f'Tableau de transposition', 0)

    return doc

def create_table_in_transposition_table(doc: docx.Document):

    n_cols = 6
    table = doc.add_table(rows=1, cols=n_cols) 
    table.style = 'Table Grid'
        
    row = table.rows[0].cells
        
    col_titles = [
        "Dispositions de la directive à transposer",
        "Normes de droit interne existantes portant déjà transposition de certaines dispositions de la directives",
        "Nature juridique des nouvelles normes à adopter pour assurerl’entière transposition de la directive",
        "Dispositions proposées",
        "Observations",
        "Assistant AI"
    ]

    for i, cell in enumerate(row):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].add_run(col_titles[i]).bold = True

    return table, n_cols   


def create_document() -> docx.Document:
    document = docx.Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    return document

def add_header(document, text):
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = text
    paragraph.style = document.styles["Header"]

def add_title(document, title):
    document.add_paragraph("")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].add_run(title).bold = True