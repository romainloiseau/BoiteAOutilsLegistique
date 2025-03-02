from io import BytesIO
import docx

from legaltoolbox.utils import plotly as myplotly
from legaltoolbox.utils import llms as llms

import streamlit as st
import numpy as np
import plotly.graph_objects as go

from hydra.utils import instantiate

from docx.enum.text import WD_ALIGN_PARAGRAPH

import legaltoolbox.utils.docx as mydocx

def compute_clustering_for_threshold(embeddings, thresh):
    labels = [0]

    n_current_embedding = 1
    for i in range(1, len(embeddings)):
        if ((embeddings[i] - embeddings[i-1]) ** 2).sum() ** 0.5 < thresh / n_current_embedding**.1:
            labels.append(labels[-1])
            n_current_embedding += 1
        else:
            labels.append(labels[-1] + 1)
            n_current_embedding = 1

    labels = np.array(labels)
    n_clusters = labels.max() + 1
    return labels, n_clusters

def run_clustering_with_order(embeddings, texts_per_clusters):
    distances = ((embeddings[1:] - embeddings[:-1])** 2).sum(axis=1) ** 0.5

    threshold = distances.max()

    k = int(len(embeddings) / texts_per_clusters)

    n_clusters_list = []
    for thresh in np.linspace(0, 10*threshold, 10000):

        labels, n_clusters = compute_clustering_for_threshold(
                embeddings, thresh)

        n_clusters_list.append(n_clusters)

        if n_clusters <= k:
            cluster_centers = np.array([embeddings[labels == i].mean(axis=0) for i in range(n_clusters)])
            break

    return labels, cluster_centers, n_clusters

@st.cache_data(show_spinner = False)
def run_cluster_and_project(_cfg, embeddings, texts, full_texts, doc_save_path):
    assert len(texts) == len(full_texts) == embeddings.shape[0]
        
    if type(full_texts) == list:
        full_texts = np.array(full_texts)

    if type(texts) == list:
        texts = np.array(texts)

    kmeans_labels, kmeans_cluster_centers, kmeans_n_clusters = run_clustering_with_order(embeddings, _cfg.tools.cluster_and_project.texts_per_clusters)
    
    chain = llms.get_chain(
        prompt = _cfg.prompts.summary.individual_cluster,
        llm = _cfg.llm,
        run_name="Clustering individual clusters"
    )
            
    clusters_description = chain.batch([{"text_to_analyse": "\n\n".join(full_texts[kmeans_labels==i])} for i in range(kmeans_n_clusters)])

    #with open(osp.join(st.session_state.logging_path, "clusters.txt"), "w") as f:
    #    f.write("\n\n".join(clusters_description))
            
    hovertexts_clusters = [f"<b>Cluster {i}</b>, {sum(kmeans_labels==i)} articles:<br>" + myplotly.format_text_for_plotly(call) for i, call in enumerate(clusters_description)]

    chain = llms.get_chain(
        prompt = _cfg.prompts.summary.report,                          
        llm = _cfg.llm,
        run_name = "Clustering reporting"
    )
            
    clustering_description = chain.invoke({
        "n_clusters": kmeans_n_clusters,
        "text_to_analyse": "\n\n".join(clusters_description)
    })

    #with open(osp.join(st.session_state.logging_path, "clusters.txt"), "w") as f:
    #    f.write(clustering_description + "\n\n\n" + "\n\n".join(clusters_description))

    tsne = instantiate(_cfg.tsne)
    projection = tsne.fit_transform(np.concatenate([embeddings, kmeans_cluster_centers]))
    projection, cluster_centers = projection[:-kmeans_n_clusters], projection[-kmeans_n_clusters:]

    hovertexts = myplotly.format_text_for_plotly(texts)
    fig = go.Figure(data=[
            go.Scatter(
                x=projection[:, 0], y=projection[:, 1],
                marker=dict(color=kmeans_labels, colorscale='Turbo', opacity=0.75),
                mode='markers', hoverinfo="text", hovertext=hovertexts, name="Articles"),
            go.Scatter(
                x=cluster_centers[:, 0], y=cluster_centers[:, 1],
                marker=dict(color=np.arange(kmeans_n_clusters), colorscale='Turbo', size=12, symbol="star"),
                mode='markers', hoverinfo="text", hovertext=hovertexts_clusters, name=f"{kmeans_n_clusters} Cluster centers")
            ]
        )
        
    fig.update_layout(title='Projection', xaxis_title=None, yaxis_title=None)
    fig.update_xaxes(showgrid=False, zeroline=False)
    fig.update_yaxes(showgrid=False, zeroline=False)
    fig.update_layout({ax:{"visible":False, "matches":None} for ax in fig.to_dict()["layout"] if "axis" in ax})

    image = BytesIO(fig.to_image(format="png", width=1000, height=1000, scale=2))

    doc, html = generate_docx_and_html(texts, full_texts, kmeans_labels, kmeans_n_clusters, clusters_description, clustering_description, image)

    #html += fig.to_html(full_html=False, include_plotlyjs='cdn')

    doc.save(doc_save_path)

    return html, fig

def generate_docx_and_html(texts, full_texts, kmeans_labels, kmeans_n_clusters, clusters_description, clustering_description, image):

    doc = mydocx.create_document()
    mydocx.add_title(doc, 'Analyse du texte')
    doc.add_paragraph(clustering_description)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    html = "<h1>Analyse du texte</h1>\n"
    html += "<p>" + clustering_description + "</p>\n"

    texts = texts.tolist()
    full_texts = full_texts.tolist()
    for i in range(kmeans_n_clusters):
        selected_articles_int = np.where(kmeans_labels == i)[0]

        if len(selected_articles_int) == 1:
            selected_articles = f"Article {selected_articles_int[0] + 1}"
        else:
            selected_articles = f"Articles {selected_articles_int[0] + 1} à {selected_articles_int[-1] + 1}"

        html += f"<h2>{selected_articles}</h2>\n"
        html += "<ul>\n"
        for j in selected_articles_int:
            html += f"<li>{texts[j]}</li>\n"
        html += "</ul>\n"
        html += "<p>" + clusters_description[i] + "</p>\n"

        doc.add_heading(f'{selected_articles}', level=1)
        for j in selected_articles_int:
            p = doc.add_paragraph(style='List Bullet')
            runner = p.add_run(texts[j])
            runner.italic = True

        for p in clusters_description[i].split("\n"):
            if p.strip() != "":
                doc.add_paragraph(p)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_picture(image, width=docx.shared.Inches(6), height=docx.shared.Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc, html