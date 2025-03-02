import streamlit as st
from tqdm.auto import tqdm

from legaltoolbox.utils import llms as llms
from legaltoolbox.utils import docx as mydocx

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os.path as osp
import logging

logger = logging.getLogger(__name__)

@st.cache_data(show_spinner = False)
def run_transposition_table(_cfg, eu_texts: list[str], _eu_text, out_analyse_impact: dict, docx_save_path: str):
    
    chain = llms.get_chain(
        prompt = _cfg.prompts.transposition_table.paragraph_analysis,
        llm = _cfg.llm,
        run_name="transposition_table"
    )

    out_ai_assistant_llm = []
    progress_bar = st.progress(0)
    for i in range(len(eu_texts)):
        if i < _cfg.tools.transposition_table.max_number_of_analysed_alineas:
            out_ai_assistant_llm.append(chain.invoke({"new_article": eu_texts[i], "context": out_analyse_impact["main_texts"] + "\n\n" + out_analyse_impact["out_ai_context"][i] if len(out_analyse_impact["out_ai_context"][i]) > 0 else out_analyse_impact["main_texts"]}))
        else:
            out_ai_assistant_llm.append("")
        progress_bar.progress((i + 1) / len(eu_texts))
    progress_bar.empty()

    doc = mydocx.create_transposition_table()
    table, n_cols = mydocx.create_table_in_transposition_table(doc)

    text_to_transpose_index = 0
    ai_assistant = []
    related_articles = []
    article = []
    color = []
    curr_article = None

    progress_bar = st.progress(0)
    for i, row in tqdm(enumerate(_eu_text.content.itertuples()), total=len(_eu_text.content)):
        table_row = table.add_row().cells

        if row.id in _eu_text.HEAD_IDS:
            for cell in range(1, n_cols):
                table_row[0].merge(table_row[cell])

            table_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table_row[0].paragraphs[0].add_run(row.text).bold = True

            ai_assistant.append(None) 
            related_articles.append(None)
            color.append(None)

            curr_article = None 
            if row.id == _eu_text.ARTICLE_TITLE_ID:
                curr_article = row.text

        elif row.id == _eu_text.CONTENT_ID:
            table_row[0].text = row.text

            if row.need_transposition != 0:
                table_row[5].paragraphs[0].add_run("Pas de transposition / " + row.need_transposition_comment).italic = True
                ai_assistant.append(None)
                related_articles.append(None)
                color.append(None)
            else:
                assert row.text == eu_texts[text_to_transpose_index]
                    
                aia = out_analyse_impact["out_ai_assistant"][text_to_transpose_index]
                for io, o in enumerate(aia):
                    mydocx.add_hyperlink(table_row[5].paragraphs[0], osp.join("https://www.legifrance.gouv.fr/codes/article_lc", o[1]), o[0])
                    if io != len(aia) - 1:
                        table_row[5].paragraphs[0].add_run("\n")
                related_articles.append(", ".join([f"{o[0]}" for o in aia]))
                    
                if len(out_ai_assistant_llm[text_to_transpose_index]) > 0:
                    if len(aia) > 0:
                        table_row[5].paragraphs[0].add_run("\n\n")

                    ai_assistant_to_add = ""

                    for ittt, ttt in enumerate(out_ai_assistant_llm[text_to_transpose_index].split("\n")):
                        if ittt == 0:
                            if "couleur" in ttt.lower():
                                if ("rouge" in ttt.lower()) or ("orange" in ttt.lower()) or ("vert" in ttt.lower()):
                                    if "rouge" in ttt.lower():
                                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="f4cccc"/>'.format(nsdecls('w')))
                                        color.append("rouge")
                                    elif "orange" in ttt.lower():
                                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="fce5cd"/>'.format(nsdecls('w')))
                                        color.append("orange")
                                    elif "vert" in ttt.lower():
                                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="fff2cc"/>'.format(nsdecls('w')))
                                        color.append("jaune")
                                    else:
                                        raise ValueError("Unknown color")
                                    table_row[5]._tc.get_or_add_tcPr().append(shading_elm_1)

                                else:
                                    table_row[5].paragraphs[0].add_run(ttt)
                                    color.append(None)
                                    ai_assistant_to_add += ttt
                                    if ittt != len(out_ai_assistant_llm[text_to_transpose_index].split("\n")) - 1:
                                        table_row[5].paragraphs[0].add_run("\n")
                            else:
                                color.append(None)
                        else:
                            table_row[5].paragraphs[0].add_run(ttt)
                            ai_assistant_to_add += "\n" + ttt
                            if ittt != len(out_ai_assistant_llm[text_to_transpose_index].split("\n")) - 1:
                                table_row[5].paragraphs[0].add_run("\n")

                    ai_assistant.append(ai_assistant_to_add)
                else:
                    ai_assistant.append(None)
                    color.append(None)
                    
            text_to_transpose_index += 1

        progress_bar.progress((i + 1) / len(_eu_text.content))
        article.append(curr_article)

    progress_bar.empty()

    doc.save(docx_save_path)

    _eu_text.content[f"Assistant AI"] = ai_assistant
    _eu_text.content[f"Articles associés"] = related_articles
    _eu_text.content["Article"] = article
    _eu_text.content["Couleur"] = color

    _eu_text.content_to_csv()

    
